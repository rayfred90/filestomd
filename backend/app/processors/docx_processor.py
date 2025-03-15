from typing import Dict, Any, BinaryIO
import docx
from .base_processor import BaseProcessor
from ..utils.chunker import DocumentChunker

class DocxProcessor(BaseProcessor):
    def __init__(self):
        super().__init__()
        self.chunker = DocumentChunker()

    def process(self, file: BinaryIO) -> Dict[str, Any]:
        """Process a DOCX file and extract its content with metadata."""
        doc = docx.Document(file)
        
        # Extract text content
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        full_text = '\n'.join(content)
        
        # Extract metadata
        metadata = self._extract_metadata(doc)
        
        # Generate chunks
        chunks = self.chunker.chunk_text(full_text, metadata)
        
        # Create markdown content
        markdown_content = self._convert_to_markdown(doc)
        
        return {
            'content': full_text,
            'markdown': markdown_content,
            'metadata': metadata,
            'chunks': chunks
        }

    def _extract_metadata(self, doc) -> Dict[str, Any]:
        """Extract metadata from the DOCX document."""
        core_properties = doc.core_properties
        
        return {
            'title': core_properties.title or '',
            'author': core_properties.author or '',
            'created': core_properties.created.isoformat() if core_properties.created else None,
            'modified': core_properties.modified.isoformat() if core_properties.modified else None,
            'word_count': len(''.join(paragraph.text for paragraph in doc.paragraphs).split()),
            'paragraph_count': len(doc.paragraphs),
            'section_count': len(doc.sections)
        }

    def _convert_to_markdown(self, doc) -> str:
        """Convert DOCX content to Markdown format."""
        markdown_lines = []
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
                
            # Handle different styles
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name[-1]  # Get heading level from style name
                markdown_lines.append(f"{'#' * int(level)} {paragraph.text}")
            else:
                # Process inline formatting
                text = paragraph.text
                for run in paragraph.runs:
                    if run.bold:
                        text = text.replace(run.text, f"**{run.text}**")
                    if run.italic:
                        text = text.replace(run.text, f"*{run.text}*")
                
                markdown_lines.append(text)
            
            # Add blank line after each paragraph
            markdown_lines.append('')
        
        return '\n'.join(markdown_lines)
